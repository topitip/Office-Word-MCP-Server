#!/usr/bin/env python3
import os
import io
import base64
import shutil
from typing import Dict, List, Optional, Any, Union, Tuple
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from mcp.server.fastmcp import FastMCP
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import sys
# Initialize FastMCP server
mcp = FastMCP("word-document-server")

# Document cache to store opened documents
documents = {}

# Helper Functions
def get_document_properties(doc_path: str, include_headers_footers: bool = False, include_notes: bool = False) -> Dict[str, Any]:
    """Get properties of a Word document."""
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        result = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
        
        # Get document sections information
        sections_info = []
        for i, section in enumerate(doc.sections):
            section_info = {
                "index": i,
                "page_width": section.page_width,
                "page_height": section.page_height,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "orientation": "portrait" if section.orientation == 0 else "landscape"
            }
            sections_info.append(section_info)
        
        result["sections"] = sections_info
        
        # Add headers and footers if requested
        if include_headers_footers:
            headers_footers = get_headers_and_footers(doc_path)
            if "error" not in headers_footers:
                result["headers_and_footers"] = headers_footers
        
        # Add footnotes and endnotes if requested
        if include_notes:
            notes = extract_footnotes_and_endnotes(doc_path)
            if "error" not in notes:
                result["notes"] = notes
        
        return result
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}

def extract_document_text(doc_path: str, include_formatting: bool = False) -> str:
    """Extract all text from a Word document.
    
    Args:
        doc_path: Path to the Word document
        include_formatting: Whether to include formatting information
        
    Returns:
        Plain text or formatted text with markup
    """
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        
        if not include_formatting:
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
                
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text.append(paragraph.text)
            
            return "\n".join(text)
        else:
            # Return text with formatting information
            formatted_text = []
            
            # Process paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                # Get paragraph style and alignment
                style_name = paragraph.style.name if paragraph.style else "Normal"
                
                # Get alignment info
                alignment = "LEFT"
                if paragraph.alignment == 1:
                    alignment = "CENTER"
                elif paragraph.alignment == 2:
                    alignment = "RIGHT"
                elif paragraph.alignment == 3:
                    alignment = "JUSTIFY"
                
                # Start paragraph with markup
                para_text = f"[PARAGRAPH style=\"{style_name}\" align=\"{alignment}\"]\n"
                
                # Process runs with formatting
                for run in paragraph.runs:
                    formatting = []
                    if run.bold:
                        formatting.append("bold")
                    if run.italic:
                        formatting.append("italic")
                    if run.underline:
                        formatting.append("underline")
                    
                    font_info = []
                    if run.font.size:
                        # Convert Pt to points
                        try:
                            size_pt = run.font.size.pt
                            font_info.append(f"size={size_pt}pt")
                        except AttributeError:
                            pass
                    
                    if run.font.name:
                        font_info.append(f"font={run.font.name}")
                    
                    if run.font.color and run.font.color.rgb:
                        font_info.append(f"color={run.font.color.rgb}")
                    
                    # Create formatting string
                    format_str = ""
                    if formatting or font_info:
                        format_str = " " + " ".join(formatting + font_info)
                    
                    # Add formatted run
                    if format_str:
                        para_text += f"[{format_str.strip()}]{run.text}[/]"
                    else:
                        para_text += run.text
                
                # End paragraph
                para_text += "\n[/PARAGRAPH]\n"
                formatted_text.append(para_text)
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_text = f"[TABLE rows={len(table.rows)} cols={len(table.columns)}]\n"
                
                for row_idx, row in enumerate(table.rows):
                    table_text += "[ROW]\n"
                    
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = "[CELL]"
                        
                        for paragraph in cell.paragraphs:
                            # Similar paragraph processing as above but simplified
                            for run in paragraph.runs:
                                if run.bold:
                                    cell_text += f"[bold]{run.text}[/bold]"
                                elif run.italic:
                                    cell_text += f"[italic]{run.text}[/italic]"
                                else:
                                    cell_text += run.text
                            
                            cell_text += "\n"
                        
                        table_text += cell_text + "[/CELL]\n"
                    
                    table_text += "[/ROW]\n"
                
                table_text += "[/TABLE]\n"
                formatted_text.append(table_text)
            
            return "".join(formatted_text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"

def get_document_styles(doc_path: str) -> Dict[str, Any]:
    """
    Get information about all styles in a document.
    
    Args:
        doc_path: Path to the Word document
        
    Returns:
        Dictionary with style information
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        styles_info = {"paragraph_styles": [], "character_styles": [], "table_styles": [], "numbering_styles": [], "other_styles": []}
        
        # Process styles
        for style in doc.styles:
            try:
                # Получаем тип стиля безопасно
                style_type = str(style.type) if hasattr(style, 'type') else "UNKNOWN"
                
                # Базовая информация о стиле, безопасная для всех типов
                style_info = {
                    "name": style.name if hasattr(style, 'name') else "Unknown",
                    "style_id": style.style_id if hasattr(style, 'style_id') else "Unknown",
                    "type": style_type
                }
                
                # Безопасно проверяем и добавляем base_style, если он существует
                # Особая осторожность для стилей нумерации
                if "NUMBERING" in style_type or "LIST" in style_type:
                    style_info["base_style"] = None
                else:
                    try:
                        if hasattr(style, 'base_style') and style.base_style:
                            style_info["base_style"] = style.base_style.name
                        else:
                            style_info["base_style"] = None
                    except AttributeError:
                        style_info["base_style"] = None
                
                # Безопасно добавляем информацию о шрифте, если она доступна
                try:
                    if hasattr(style, 'font'):
                        font_info = {}
                        if hasattr(style.font, 'name'):
                            font_info["name"] = style.font.name
                        if hasattr(style.font, 'size') and style.font.size:
                            try:
                                font_info["size"] = style.font.size.pt
                            except AttributeError:
                                pass
                        if hasattr(style.font, 'bold'):
                            font_info["bold"] = style.font.bold
                        if hasattr(style.font, 'italic'):
                            font_info["italic"] = style.font.italic
                        if hasattr(style.font, 'underline'):
                            font_info["underline"] = style.font.underline
                        if hasattr(style.font, 'color') and style.font.color and hasattr(style.font.color, 'rgb'):
                            font_info["color"] = style.font.color.rgb
                        
                        if font_info:  # Добавляем только если есть информация
                            style_info["font"] = font_info
                except Exception:
                    pass  # Игнорируем ошибки доступа к свойствам шрифта
                
                # Безопасно добавляем информацию о формате параграфа, если она доступна
                try:
                    if hasattr(style, 'paragraph_format'):
                        para_format = {}
                        if hasattr(style.paragraph_format, 'alignment'):
                            para_format["alignment"] = str(style.paragraph_format.alignment) if style.paragraph_format.alignment else "LEFT"
                        if hasattr(style.paragraph_format, 'left_indent'):
                            para_format["left_indent"] = style.paragraph_format.left_indent
                        if hasattr(style.paragraph_format, 'right_indent'):
                            para_format["right_indent"] = style.paragraph_format.right_indent
                        if hasattr(style.paragraph_format, 'first_line_indent'):
                            para_format["first_line_indent"] = style.paragraph_format.first_line_indent
                        if hasattr(style.paragraph_format, 'line_spacing'):
                            para_format["line_spacing"] = style.paragraph_format.line_spacing
                        if hasattr(style.paragraph_format, 'space_before'):
                            para_format["space_before"] = style.paragraph_format.space_before
                        if hasattr(style.paragraph_format, 'space_after'):
                            para_format["space_after"] = style.paragraph_format.space_after
                        
                        if para_format:  # Добавляем только если есть информация
                            style_info["paragraph_format"] = para_format
                except Exception:
                    pass  # Игнорируем ошибки доступа к свойствам формата параграфа
                
                # Определяем тип стиля и добавляем в соответствующий список
                if "PARAGRAPH" in style_type:
                    styles_info["paragraph_styles"].append(style_info)
                elif "CHARACTER" in style_type:
                    styles_info["character_styles"].append(style_info)
                elif "TABLE" in style_type:
                    styles_info["table_styles"].append(style_info)
                elif "NUMBERING" in style_type or "LIST" in style_type:
                    styles_info["numbering_styles"].append(style_info)
                else:
                    styles_info["other_styles"].append(style_info)
            except Exception as style_error:
                # Если с конкретным стилем возникла проблема, добавим его в список с ошибкой,
                # но не прервем обработку всех стилей
                error_style = {
                    "name": getattr(style, "name", "Unknown"),
                    "error": str(style_error)
                }
                styles_info["other_styles"].append(error_style)
        
        return styles_info
    except Exception as e:
        return {"error": f"Failed to get document styles: {str(e)}"}

def get_table_detailed_info(table, table_index: int) -> Dict[str, Any]:
    """
    Get detailed information about a table including cell formatting.
    
    Args:
        table: docx Table object
        table_index: Index of the table in the document
        
    Returns:
        Dictionary with detailed table information
    """
    table_info = {
        "index": table_index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "style": table.style.name if hasattr(table, 'style') and table.style else "None",
        "alignment": str(table.alignment) if hasattr(table, 'alignment') and table.alignment else "LEFT",
        "cells": []
    }
    
    # Get cell information
    for i, row in enumerate(table.rows):
        row_cells = []
        for j, cell in enumerate(row.cells):
            # Get basic cell information
            cell_info = {
                "row": i,
                "column": j,
                "text": cell.text,
                "paragraphs": []
            }
            
            # Get cell border information if available
            try:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = tcPr.first_child_found_in("w:tcBorders")
                if borders is not None:
                    cell_info["borders"] = {}
                    for border in ["top", "bottom", "left", "right"]:
                        b = borders.find(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{border}")
                        if b is not None:
                            cell_info["borders"][border] = {
                                "val": b.get(qn("w:val")),
                                "color": b.get(qn("w:color")),
                                "size": b.get(qn("w:sz"))
                            }
            except:
                pass
            
            # Get cell background shading if available
            try:
                shading = tcPr.first_child_found_in("w:shd")
                if shading is not None:
                    fill = shading.get(qn("w:fill"))
                    if fill:
                        cell_info["shading"] = fill
            except:
                pass
            
            # Get paragraph formatting in the cell
            for p_idx, paragraph in enumerate(cell.paragraphs):
                para_info = {
                    "text": paragraph.text,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "alignment": str(paragraph.alignment) if paragraph.alignment else "LEFT",
                    "runs": []
                }
                
                # Get run formatting
                for run in paragraph.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline
                    }
                    
                    if run.font.size:
                        try:
                            run_info["font_size"] = run.font.size.pt
                        except:
                            pass
                    
                    para_info["runs"].append(run_info)
                
                cell_info["paragraphs"].append(para_info)
            
            row_cells.append(cell_info)
        
        table_info["cells"].append(row_cells)
    
    return table_info

def get_document_structure(doc_path: str, detailed_tables: bool = False) -> Dict[str, Any]:
    """Get the structure of a Word document.
    
    Args:
        doc_path: Path to the Word document
        detailed_tables: Whether to include detailed table information
        
    Returns:
        Document structure dictionary
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            # Определение выравнивания
            alignment = "LEFT"
            if para.alignment == 1:
                alignment = "CENTER"
            elif para.alignment == 2:
                alignment = "RIGHT"
            elif para.alignment == 3:
                alignment = "JUSTIFY"
            
            # Получение информации о форматировании
            format_info = {
                "indent_left": para.paragraph_format.left_indent,
                "indent_right": para.paragraph_format.right_indent,
                "indent_first_line": para.paragraph_format.first_line_indent,
                "space_before": para.paragraph_format.space_before,
                "space_after": para.paragraph_format.space_after,
                "line_spacing": para.paragraph_format.line_spacing,
            }
            
            # Информация о параграфе
            paragraph_info = {
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal",
                "alignment": alignment,
                "format": format_info
            }
            
            # Получение информации о форматировании текста
            if para.runs:
                runs_info = []
                for run in para.runs:
                    runs_info.append({
                        "text": run.text[:50] + ("..." if len(run.text) > 50 else ""),
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_size": run.font.size,
                        "font_name": run.font.name,
                        "highlight_color": run.font.highlight_color,
                        "color": run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    })
                paragraph_info["runs"] = runs_info
            
            structure["paragraphs"].append(paragraph_info)
        
        # Get tables
        for i, table in enumerate(doc.tables):
            if detailed_tables:
                # Get detailed table information
                table_info = get_table_detailed_info(table, i)
                structure["tables"].append(table_info)
            else:
                # Get basic table information
                table_data = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "preview": []
                }
                
                # Get sample of table data
                max_rows = min(3, len(table.rows))
                for row_idx in range(max_rows):
                    row_data = []
                    max_cols = min(3, len(table.columns))
                    for col_idx in range(max_cols):
                        try:
                            cell_text = table.cell(row_idx, col_idx).text
                            row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                        except IndexError:
                            row_data.append("N/A")
                    table_data["preview"].append(row_data)
                
                structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}

def check_file_writeable(filepath: str) -> Tuple[bool, str]:
    """
    Check if a file can be written to.
    
    Args:
        filepath: Path to the file
        
    Returns:
        Tuple of (is_writeable, error_message)
    """
    # If file doesn't exist, check if directory is writeable
    if not os.path.exists(filepath):
        directory = os.path.dirname(filepath)
        if not os.path.exists(directory):
            return False, f"Directory {directory} does not exist"
        if not os.access(directory, os.W_OK):
            return False, f"Directory {directory} is not writeable"
        return True, ""
    
    # If file exists, check if it's writeable
    if not os.access(filepath, os.W_OK):
        return False, f"File {filepath} is not writeable (permission denied)"
    
    # Try to open the file for writing to see if it's locked
    try:
        with open(filepath, 'a'):
            pass
        return True, ""
    except IOError as e:
        return False, f"File {filepath} is not writeable: {str(e)}"
    except Exception as e:
        return False, f"Unknown error checking file permissions: {str(e)}"

def create_document_copy(source_path: str, dest_path = None) -> Tuple[bool, str, Optional[str]]:
    """
    Create a copy of a document.
    
    Args:
        source_path: Path to the source document
        dest_path: Optional path for the new document. If not provided, will use source_path + '_copy.docx'
        
    Returns:
        Tuple of (success, message, new_filepath)
    """
    if not os.path.exists(source_path):
        return False, f"Source document {source_path} does not exist", None
    
    if not dest_path:
        # Generate a new filename if not provided
        base, ext = os.path.splitext(source_path)
        dest_path = f"{base}_copy{ext}"
    
    try:
        # Simple file copy
        shutil.copy2(source_path, dest_path)
        return True, f"Document copied to {dest_path}", dest_path
    except Exception as e:
        return False, f"Failed to copy document: {str(e)}", None

def ensure_heading_style(doc):
    """
    Ensure Heading styles exist in the document.
    
    Args:
        doc: Document object
    """
    for i in range(1, 10):  # Create Heading 1 through Heading 9
        style_name = f'Heading {i}'
        try:
            # Try to access the style to see if it exists
            style = doc.styles[style_name]
        except KeyError:
            # Create the style if it doesn't exist
            try:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                if i == 1:
                    style.font.size = Pt(16)
                    style.font.bold = True
                elif i == 2:
                    style.font.size = Pt(14)
                    style.font.bold = True
                else:
                    style.font.size = Pt(12)
                    style.font.bold = True
            except Exception:
                # If style creation fails, we'll just use default formatting
                pass

def ensure_table_style(doc):
    """
    Ensure Table Grid style exists in the document.
    
    Args:
        doc: Document object
    """
    try:
        # Try to access the style to see if it exists
        style = doc.styles['Table Grid']
    except KeyError:
        # If style doesn't exist, we'll handle it at usage time
        pass

# MCP Tools
@mcp.tool(name="create_document")
async def create_document(filename: str, title = None, author = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"

@mcp.tool(name="add_heading")
async def add_heading(filename: str, text: str, level: int = 1, alignment = None) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        alignment: Optional alignment ('left', 'center', 'right', 'justify')
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        
        # Ensure heading styles exist
        ensure_heading_style(doc)
        
        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
            
            # Set alignment if specified
            if alignment:
                alignment_map = {
                    'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                    'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                    'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    heading.alignment = alignment_map[alignment.lower()]
            
            doc.save(filename)
            return f"Heading '{text}' (level {level}) added to {filename}"
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            paragraph = doc.add_paragraph(text)
            paragraph.style = doc.styles['Normal']
            run = paragraph.runs[0]
            run.bold = True
            # Adjust size based on heading level
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            
            # Set alignment if specified
            if alignment:
                alignment_map = {
                    'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                    'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                    'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    paragraph.alignment = alignment_map[alignment.lower()]
            
            doc.save(filename)
            return f"Heading '{text}' added to {filename} with direct formatting (style not available)"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"

@mcp.tool(name="add_paragraph")
async def add_paragraph(filename: str, text: str, style = None, alignment = None) -> str:
    """Add a paragraph to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        alignment: Optional alignment ('left', 'center', 'right', 'justify')
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)
        
        # Применяем стиль, если указан
        if style:
            try:
                # Пробуем применить стиль напрямую
                try:
                    paragraph.style = doc.styles[style]
                except KeyError:
                    # Если стиль не найден по имени, ищем по id
                    style_found = False
                    for s in doc.styles:
                        if s.name.lower() == style.lower() or (hasattr(s, 'style_id') and s.style_id.lower() == style.lower()):
                            paragraph.style = s
                            style_found = True
                            break
                    
                    if not style_found:
                        # Если стиль не найден, используем Normal и сообщаем об этом
                        paragraph.style = doc.styles['Normal']
                        doc.save(filename)
                        return f"Style '{style}' not found, paragraph added with default style to {filename}"
            except Exception as style_error:
                # В случае ошибки применения стиля
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Error applying style '{style}': {str(style_error)}. Paragraph added with default style to {filename}"
        
        # Set alignment if specified
        if alignment:
            alignment_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            if alignment.lower() in alignment_map:
                paragraph.alignment = alignment_map[alignment.lower()]
        
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"

@mcp.tool(name="add_table")
async def add_table(filename: str, rows: int, cols: int, data = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            # This is a simplified approach - complete border styling would require more code
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"

@mcp.tool(name="add_picture")
async def add_picture(filename: str, image_path: str, width = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"

@mcp.tool(name="get_document_info")
async def get_document_info(filename: str, include_headers_footers: bool = False, include_notes: bool = False) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
        include_headers_footers: Whether to include headers and footers information
        include_notes: Whether to include footnotes and endnotes information
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename, include_headers_footers, include_notes)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"

@mcp.tool(name="get_document_text")
async def get_document_text(filename: str, include_formatting: bool = False) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
        include_formatting: Whether to include formatting information
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    return extract_document_text(filename, include_formatting)

@mcp.tool(name="get_document_styles")
async def get_document_styles_tool(filename: str) -> str:
    """Get information about all styles in a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        styles_info = get_document_styles(filename)
        return json.dumps(styles_info, indent=2)
    except Exception as e:
        return f"Failed to get document styles: {str(e)}"

@mcp.tool(name="get_document_outline")
async def get_document_outline(filename: str, detailed_tables: bool = False) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
        detailed_tables: Whether to include detailed table information
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    structure = get_document_structure(filename, detailed_tables)
    return json.dumps(structure, indent=2)

@mcp.tool(name="list_available_documents")
async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"

@mcp.tool(name="copy_document")
async def copy_document(source_filename: str, destination_filename = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    if not source_filename.endswith('.docx'):
        source_filename += '.docx'
    
    if destination_filename and not destination_filename.endswith('.docx'):
        destination_filename += '.docx'
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"

# Resources
@mcp.resource("docx:{path}")
async def document_resource(path: str) -> str:
    """Access Word document content."""
    if not path.endswith('.docx'):
        path += '.docx'
    
    if not os.path.exists(path):
        return f"Document {path} does not exist"
    
    return extract_document_text(path, include_formatting=False)

@mcp.resource("docx-formatted:{path}")
async def formatted_document_resource(path: str) -> str:
    """Access Word document content with formatting information."""
    if not path.endswith('.docx'):
        path += '.docx'
    
    if not os.path.exists(path):
        return f"Document {path} does not exist"
    
    return extract_document_text(path, include_formatting=True)

def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs

def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count

def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    
    Args:
        cell: The cell to modify
        **kwargs: Border properties (top, bottom, left, right, val, color)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for key, value in kwargs.items():
        if key in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(key)
            
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs.get('val', 'single'))
            element.set(qn('w:sz'), kwargs.get('sz', '4'))
            element.set(qn('w:space'), kwargs.get('space', '0'))
            element.set(qn('w:color'), kwargs.get('color', 'auto'))
            
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            tcBorders.append(element)

def create_style(doc, style_name, style_type, base_style=None, font_properties=None, paragraph_properties=None):
    """
    Create a new style in the document.
    
    Args:
        doc: Document object
        style_name: Name for the new style
        style_type: Type of style (WD_STYLE_TYPE)
        base_style: Optional base style to inherit from
        font_properties: Dictionary of font properties (bold, italic, size, name, color)
        paragraph_properties: Dictionary of paragraph properties (alignment, spacing)
        
    Returns:
        The created style
    """
    try:
        # Проверяем, существует ли стиль по имени
        try:
            style = doc.styles[style_name]
            return style
        except KeyError:
            # Стиль не найден, создаем новый
            new_style = doc.styles.add_style(style_name, style_type)
            
            # Set base style if specified
            if base_style:
                try:
                    new_style.base_style = doc.styles[base_style]
                except KeyError:
                    # Если базовый стиль не найден, используем Normal
                    new_style.base_style = doc.styles['Normal']
            
            # Set font properties
            if font_properties:
                font = new_style.font
                if 'bold' in font_properties:
                    font.bold = font_properties['bold']
                if 'italic' in font_properties:
                    font.italic = font_properties['italic']
                if 'size' in font_properties:
                    font.size = Pt(font_properties['size'])
                if 'name' in font_properties:
                    font.name = font_properties['name']
                if 'color' in font_properties:
                    try:
                        # Обработка цветов
                        color_value = font_properties['color']
                        # Для известных цветов используем RGB значения
                        color_map = {
                            'red': RGBColor(255, 0, 0),
                            'blue': RGBColor(0, 0, 255),
                            'green': RGBColor(0, 128, 0),
                            'yellow': RGBColor(255, 255, 0),
                            'black': RGBColor(0, 0, 0),
                            'white': RGBColor(255, 255, 255),
                        }
                        
                        if isinstance(color_value, str) and color_value.lower() in color_map:
                            font.color.rgb = color_map[color_value.lower()]
                        else:
                            # Для других случаев пробуем установить прямо
                            font.color.rgb = color_value
                    except Exception:
                        # В случае ошибки игнорируем установку цвета
                        pass
            
            # Set paragraph properties
            if paragraph_properties:
                if 'alignment' in paragraph_properties:
                    new_style.paragraph_format.alignment = paragraph_properties['alignment']
                if 'spacing' in paragraph_properties:
                    new_style.paragraph_format.line_spacing = paragraph_properties['spacing']
            
            return new_style
    except Exception as e:
        print(f"Error creating style: {e}")
        # Возвращаем None в случае ошибки
        return None

# Add these MCP tools to the existing set

@mcp.tool(name="format_text")
async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                     bold = None, italic = None, 
                     underline = None, color = None,
                     font_size = None, font_name = None) -> str:
    """Format a specific range of text within a paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        start_pos: Start position within the paragraph text
        end_pos: End position within the paragraph text
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Преобразование строковых параметров в числовые
        try:
            paragraph_index = int(paragraph_index)
            start_pos = int(start_pos)
            end_pos = int(end_pos)
        except (ValueError, TypeError):
            return "Position parameters must be integers"
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
            return f"Invalid text positions. Paragraph has {len(text)} characters."
        
        # Get the text to format
        target_text = text[start_pos:end_pos]
        
        # Clear existing runs and create three runs: before, target, after
        for run in paragraph.runs:
            run.clear()
        
        # Add text before target
        if start_pos > 0:
            run_before = paragraph.add_run(text[:start_pos])
        
        # Add target text with formatting
        run_target = paragraph.add_run(target_text)
        
        # Преобразуем строковые значения в bool, если необходимо
        if bold is not None:
            if isinstance(bold, str):
                if bold.lower() == 'true':
                    run_target.bold = True
                elif bold.lower() == 'false':
                    run_target.bold = False
            else:
                run_target.bold = bool(bold)
                
        if italic is not None:
            if isinstance(italic, str):
                if italic.lower() == 'true':
                    run_target.italic = True
                elif italic.lower() == 'false':
                    run_target.italic = False
            else:
                run_target.italic = bool(italic)
                
        if underline is not None:
            if isinstance(underline, str):
                if underline.lower() == 'true':
                    run_target.underline = True
                elif underline.lower() == 'false':
                    run_target.underline = False
            else:
                run_target.underline = bool(underline)
                
        if color:
            try:
                # Карта цветов для распространенных имен цветов
                color_map = {
                    'red': RGBColor(255, 0, 0),
                    'blue': RGBColor(0, 0, 255),
                    'green': RGBColor(0, 128, 0),
                    'yellow': RGBColor(255, 255, 0),
                    'black': RGBColor(0, 0, 0),
                    'white': RGBColor(255, 255, 255),
                }
                
                if isinstance(color, str) and color.lower() in color_map:
                    run_target.font.color.rgb = color_map[color.lower()]
                else:
                    # Пробуем установить цвет по индексу
                    try:
                        index_color_map = {
                            'red': WD_COLOR_INDEX.RED,
                            'blue': WD_COLOR_INDEX.BLUE,
                            'green': WD_COLOR_INDEX.GREEN,
                            'yellow': WD_COLOR_INDEX.YELLOW,
                            'black': WD_COLOR_INDEX.BLACK,
                        }
                        if isinstance(color, str) and color.lower() in index_color_map:
                            run_target.font.color.index = index_color_map[color.lower()]
                    except Exception:
                        # Если не сработало, игнорируем
                        pass
            except Exception:
                # В случае ошибки игнорируем установку цвета
                pass
                
        if font_size:
            try:
                font_size_value = int(font_size)
                run_target.font.size = Pt(font_size_value)
            except (ValueError, TypeError):
                # Игнорируем ошибки преобразования размера шрифта
                pass
                
        if font_name:
            run_target.font.name = str(font_name)
        
        # Add text after target
        if end_pos < len(text):
            run_after = paragraph.add_run(text[end_pos:])
        
        doc.save(filename)
        return f"Text '{target_text}' formatted successfully in paragraph {paragraph_index}."
    except Exception as e:
        return f"Failed to format text: {str(e)}"

@mcp.tool(name="search_and_replace")
async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

@mcp.tool(name="delete_paragraph")
async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"

@mcp.tool(name="create_custom_style")
async def create_custom_style(filename: str, style_name: str, 
                             bold = None, italic = None,
                             font_size = None, font_name = None,
                             color = None, base_style = None) -> str:
    """Create a custom style in the document.
    
    Args:
        filename: Path to the Word document
        style_name: Name for the new style
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        font_size: Font size in points
        font_name: Font name/family
        color: Text color (e.g., 'red', 'blue')
        base_style: Optional existing style to base this on
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Build font properties dictionary
        font_properties = {}
        
        # Преобразуем строковые значения в bool, если необходимо
        if bold is not None:
            if isinstance(bold, str):
                if bold.lower() == 'true':
                    font_properties['bold'] = True
                elif bold.lower() == 'false':
                    font_properties['bold'] = False
            else:
                font_properties['bold'] = bool(bold)
                
        if italic is not None:
            if isinstance(italic, str):
                if italic.lower() == 'true':
                    font_properties['italic'] = True
                elif italic.lower() == 'false':
                    font_properties['italic'] = False
            else:
                font_properties['italic'] = bool(italic)
        
        if font_size is not None:
            try:
                font_properties['size'] = int(font_size)
            except (ValueError, TypeError):
                pass
        
        if font_name is not None:
            font_properties['name'] = str(font_name)
            
        if color is not None:
            font_properties['color'] = str(color)
        
        # Create the style
        new_style = create_style(
            doc, 
            style_name, 
            WD_STYLE_TYPE.PARAGRAPH, 
            base_style=base_style,
            font_properties=font_properties
        )
        
        doc.save(filename)
        return f"Style '{style_name}' created successfully."
    except Exception as e:
        return f"Failed to create style: {str(e)}"

@mcp.tool(name="format_table")
async def format_table(filename: str, table_index: int, 
                      has_header_row = None,
                      border_style = None,
                      shading = None) -> str:
    """Format a table with borders, shading, and structure.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        try:
            table_index = int(table_index)
        except (ValueError, TypeError):
            return "Table index must be an integer"
            
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Format header row if requested
        if has_header_row:
            # Преобразуем строковое значение в bool, если необходимо
            is_header = False
            if isinstance(has_header_row, str):
                if has_header_row.lower() == 'true':
                    is_header = True
            else:
                is_header = bool(has_header_row)
                
            if is_header and table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Apply border style if specified
        if border_style:
            val_map = {
                'none': 'nil',
                'single': 'single',
                'double': 'double',
                'thick': 'thick'
            }
            val = val_map.get(str(border_style).lower(), 'single')
            
            # Apply to all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top=True,
                        bottom=True,
                        left=True,
                        right=True,
                        val=val,
                        color="000000"
                    )
        
        # Apply cell shading if specified
        if shading:
            # Преобразуем строковое значение JSON в список, если необходимо
            if isinstance(shading, str):
                try:
                    import json
                    shading = json.loads(shading)
                except json.JSONDecodeError:
                    return "Invalid shading format. Expected 2D array as JSON string."
            
            for i, row_colors in enumerate(shading):
                if i >= len(table.rows):
                    break
                for j, color in enumerate(row_colors):
                    if j >= len(table.rows[i].cells):
                        break
                    try:
                        # Apply shading to cell
                        cell = table.rows[i].cells[j]
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        # Skip if color format is invalid
                        pass
        
        doc.save(filename)
        return f"Table at index {table_index} formatted successfully."
    except Exception as e:
        return f"Failed to format table: {str(e)}"

@mcp.tool(name="add_page_break")
async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"

@mcp.tool(name="set_paragraph_alignment")
async def set_paragraph_alignment(filename: str, paragraph_index: int, alignment: str) -> str:
    """Set the alignment for a paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        alignment: Alignment to set ('left', 'center', 'right', 'justify')
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Validate alignment
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        
        if alignment.lower() not in alignment_map:
            return f"Invalid alignment. Supported values: left, center, right, justify."
        
        # Set alignment
        paragraph = doc.paragraphs[paragraph_index]
        paragraph.alignment = alignment_map[alignment.lower()]
        
        doc.save(filename)
        return f"Alignment for paragraph {paragraph_index} set to '{alignment}'."
    except Exception as e:
        return f"Failed to set paragraph alignment: {str(e)}"

def get_headers_and_footers(doc_path: str) -> Dict[str, Any]:
    """
    Get header and footer information from a Word document.
    
    Args:
        doc_path: Path to the Word document
        
    Returns:
        Dictionary with header and footer information
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        result = {
            "headers": [],
            "footers": []
        }
        
        # Process each section
        for i, section in enumerate(doc.sections):
            section_info = {"section_index": i}
            
            # Process headers
            header_info = {}
            if section.header.is_linked_to_previous:
                header_info["linked_to_previous"] = True
            else:
                header_info["linked_to_previous"] = False
                header_text = []
                for paragraph in section.header.paragraphs:
                    header_text.append(paragraph.text)
                header_info["text"] = "\n".join(header_text)
                
                # Add formatting information
                runs_with_formatting = []
                for para in section.header.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run_info = {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline
                            }
                            runs_with_formatting.append(run_info)
                
                header_info["formatted_runs"] = runs_with_formatting
            
            section_info["header"] = header_info
            result["headers"].append(section_info)
            
            # Process footers
            footer_info = {}
            if section.footer.is_linked_to_previous:
                footer_info["linked_to_previous"] = True
            else:
                footer_info["linked_to_previous"] = False
                footer_text = []
                for paragraph in section.footer.paragraphs:
                    footer_text.append(paragraph.text)
                footer_info["text"] = "\n".join(footer_text)
                
                # Add formatting information
                runs_with_formatting = []
                for para in section.footer.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run_info = {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline
                            }
                            runs_with_formatting.append(run_info)
                
                footer_info["formatted_runs"] = runs_with_formatting
            
            section_info["footer"] = footer_info
            result["footers"].append(section_info)
        
        return result
    except Exception as e:
        return {"error": f"Failed to get headers and footers: {str(e)}"}

def extract_footnotes_and_endnotes(doc_path: str) -> Dict[str, Any]:
    """
    Extract footnotes and endnotes from a Word document.
    Note: This is limited by python-docx capabilities and accesses the underlying XML.
    
    Args:
        doc_path: Path to the Word document
        
    Returns:
        Dictionary with footnote and endnote information
    """
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        result = {
            "footnotes": [],
            "endnotes": []
        }
        
        # Get access to the document part
        document_part = doc._part
        
        # Try to get footnotes
        try:
            if hasattr(document_part, 'footnotes_part') and document_part.footnotes_part:
                footnotes_part = document_part.footnotes_part
                # Access the footnotes XML
                footnotes_xml = footnotes_part._element
                
                # Extract footnotes - this is a simplified approach
                for footnote in footnotes_xml.findall('.//w:footnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    footnote_id = footnote.get(qn('w:id'))
                    if footnote_id and footnote_id not in ('-1', '0'):  # Skip special footnotes
                        footnote_text = ""
                        for paragraph in footnote.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for text_element in paragraph.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                footnote_text += text_element.text if text_element.text else ""
                        
                        result["footnotes"].append({
                            "id": footnote_id,
                            "text": footnote_text
                        })
        except Exception as footnote_error:
            result["footnote_error"] = str(footnote_error)
        
        # Try to get endnotes
        try:
            if hasattr(document_part, 'endnotes_part') and document_part.endnotes_part:
                endnotes_part = document_part.endnotes_part
                # Access the endnotes XML
                endnotes_xml = endnotes_part._element
                
                # Extract endnotes - this is a simplified approach
                for endnote in endnotes_xml.findall('.//w:endnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    endnote_id = endnote.get(qn('w:id'))
                    if endnote_id and endnote_id not in ('-1', '0'):  # Skip special endnotes
                        endnote_text = ""
                        for paragraph in endnote.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for text_element in paragraph.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                endnote_text += text_element.text if text_element.text else ""
                        
                        result["endnotes"].append({
                            "id": endnote_id,
                            "text": endnote_text
                        })
        except Exception as endnote_error:
            result["endnote_error"] = str(endnote_error)
        
        return result
    except Exception as e:
        return {"error": f"Failed to extract footnotes and endnotes: {str(e)}"}

@mcp.tool(name="get_headers_and_footers")
async def get_headers_and_footers_tool(filename: str) -> str:
    """Get header and footer information from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        result = get_headers_and_footers(filename)
        return json.dumps(result, indent=2)
    except Exception as e:
        return f"Failed to get headers and footers: {str(e)}"

@mcp.tool(name="get_footnotes_and_endnotes")
async def get_footnotes_and_endnotes_tool(filename: str) -> str:
    """Extract footnotes and endnotes from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        result = extract_footnotes_and_endnotes(filename)
        return json.dumps(result, indent=2)
    except Exception as e:
        return f"Failed to extract notes: {str(e)}"

# Main execution point
def main():
    """Entry point for the MCP server."""
    # Run the server
    mcp.run(transport='stdio')

if __name__ == "__main__":
    main()